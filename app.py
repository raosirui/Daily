from flask import Flask, render_template, request, redirect, url_for, session, jsonify, send_file
from flask_session import Session
import json
import os
import datetime
import re
import logging
from docx import Document
from docx.shared import Inches
from werkzeug.utils import secure_filename

app = Flask(__name__)

# 为Jinja2环境添加Python内置函数支持
app.jinja_env.globals.update(max=max)
app.secret_key = 'your-secret-key'
app.config['SESSION_TYPE'] = 'filesystem'
app.config['SESSION_FILE_DIR'] = './flask_session'
Session(app)

# 日志配置
if not os.path.exists('logs'):
    os.makedirs('logs')

log_filename = os.path.join('logs', datetime.datetime.now().strftime('%Y%m%d') + '.log')
logging.basicConfig(
    filename=log_filename,
    level=logging.INFO,
    format='%(asctime)s - %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)

# 工具函数

def ensure_directories():
    """确保必要的目录结构存在"""
    directories = ['data/daily', 'data/docx']
    for directory in directories:
        if not os.path.exists(directory):
            os.makedirs(directory, exist_ok=True)
def log_activity(username, activity):
    """记录用户活动"""
    log_filename = os.path.join('logs', datetime.datetime.now().strftime('%Y%m%d') + '.log')
    with open(log_filename, 'a', encoding='utf-8') as f:
        f.write(f"{datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')} - {username} - {activity}\n")

def load_passwords():
    """加载用户密码"""
    passwords = {}
    try:
        with open('passwd.txt', 'r', encoding='utf-8') as f:
            for line in f:
                if line.strip():
                    username, password = line.strip().split(':', 1)
                    passwords[username] = password
    except Exception as e:
        print(f"Error loading passwords: {e}")
    return passwords

def load_projects():
    """加载项目列表"""
    try:
        if os.path.exists('project_list.json'):
            with open('project_list.json', 'r', encoding='utf-8') as f:
                return json.load(f)
    except Exception as e:
        print(f"Error loading projects: {e}")
    return []

def save_projects(projects):
    """保存项目列表"""
    try:
        with open('project_list.json', 'w', encoding='utf-8') as f:
            json.dump(projects, f, ensure_ascii=False, indent=2)
        return True
    except Exception as e:
        print(f"Error saving projects: {e}")
        return False

def save_daily_report(username, date, report_data):
    """保存日报数据"""
    # 确保必要的目录存在
    ensure_directories()
    # 确保使用统一的日期格式(YYYYMMDD)，移除可能存在的连字符
    normalized_date = date.replace('-', '')
    filename = f"data/daily/{normalized_date}_{username}.json"
    try:
        with open(filename, 'w', encoding='utf-8') as f:
            json.dump(report_data, f, ensure_ascii=False, indent=2)
        return True
    except Exception as e:
        print(f"Error saving daily report: {e}")
        return False

def load_daily_report(username, date):
    """加载日报数据"""
    # 确保必要的目录存在
    ensure_directories()
    # 确保使用统一的日期格式(YYYYMMDD)，移除可能存在的连字符
    normalized_date = date.replace('-', '')
    filename = f"data/daily/{normalized_date}_{username}.json"
    try:
        if os.path.exists(filename):
            with open(filename, 'r', encoding='utf-8') as f:
                return json.load(f)
    except Exception as e:
        print(f"Error loading daily report: {e}")
    return []

def generate_docx_report(date, data):
    """生成Word文档报告"""
    # 确保必要的目录存在
    ensure_directories()
    
    doc = Document()
    
    # 获取星期几
    date_obj = datetime.datetime.strptime(date, '%Y%m%d')
    weekday = date_obj.strftime('%A')
    weekday_map = {
        'Monday': '星期一',
        'Tuesday': '星期二',
        'Wednesday': '星期三',
        'Thursday': '星期四',
        'Friday': '星期五',
        'Saturday': '星期六',
        'Sunday': '星期日'
    }
    weekday_cn = weekday_map.get(weekday, weekday)
    
    # 设置文档样式
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.shared import OxmlElement, qn
    
    # 添加标题并设置为黑体
    title = doc.add_heading('', level=0)
    title_run = title.add_run(f'市场部工作日志{date_obj.year}年{date_obj.month}月{date_obj.day}日（{weekday_cn}）')
    title_run.font.name = '黑体'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    title_run.font.size = Pt(16)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加表格并设置列宽
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # 设置表格列宽
    table.cell(0, 0).width = Inches(1.2)
    table.cell(0, 1).width = Inches(3.5)
    table.cell(0, 2).width = Inches(3.5)
    
    # 设置表头
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '名称'
    hdr_cells[1].text = '当日工作计划'
    hdr_cells[2].text = '次日工作计划'
    
    # 设置表头字体为黑体并居中
    for cell in hdr_cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            run.font.bold = True
    
    # 填充表格数据
    for user, works in data.items():
        row_cells = table.add_row().cells
        
        # 设置用户名单元格
        user_cell = row_cells[0]
        user_cell.text = user
        user_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in user_cell.paragraphs[0].runs:
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        
        # 设置当日工作计划单元格
        today_cell = row_cells[1]
        today_cell.text = works.get('d-1', '无')
        for run in today_cell.paragraphs[0].runs:
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        today_cell.paragraphs[0].space_after = Pt(6)
        
        # 设置次日工作计划单元格
        tomorrow_cell = row_cells[2]
        tomorrow_cell.text = works.get('d', '无')
        for run in tomorrow_cell.paragraphs[0].runs:
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        tomorrow_cell.paragraphs[0].space_after = Pt(6)
    
    # 保存文档到data/docx目录
    filename = f"市场部工作日志{date_obj.year}.{date_obj.month:02d}.{date_obj.day:02d}.docx"
    full_path = f"data/docx/{filename}"
    doc.save(full_path)
    return full_path

def get_all_users():
    """获取所有用户列表"""
    passwords = load_passwords()
    return list(passwords.keys())

def search_projects(keyword):
    """搜索项目，支持空格分隔的AND匹配，使用正则表达式*<搜索内容>*匹配"""
    projects = load_projects()
    # 添加空选项
    results = ['']  # 空选项，允许不关联项目
    
    if not keyword:
        # 如果没有搜索词，只返回空选项和所有项目
        return results + projects
    
    keywords = keyword.split()
    
    for project in projects:
        match = True
        for k in keywords:
            # 使用正则表达式进行*<搜索内容>*匹配
            # 转换为不区分大小写的匹配
            if not re.search(re.escape(k), project, re.IGNORECASE):
                match = False
                break
        if match:
            results.append(project)
    
    return results

# 路由
@app.route('/')
def index():
    if 'username' in session:
        return redirect(url_for('daily_report'))
    return render_template('login.html')

@app.route('/login', methods=['POST'])
def login():
    username = request.form.get('username')
    password = request.form.get('password')
    
    passwords = load_passwords()
    if username in passwords and passwords[username] == password:
        session['username'] = username
        log_activity(username, 'Login')
        return redirect(url_for('daily_report'))
    else:
        return render_template('login.html', error='用户名或密码错误')

@app.route('/logout')
def logout():
    if 'username' in session:
        log_activity(session['username'], 'Logout')
        session.pop('username', None)
    return redirect(url_for('index'))

@app.route('/daily_report')
def daily_report():
    if 'username' not in session:
        return redirect(url_for('index'))
    
    username = session['username']
    today = datetime.datetime.now().strftime('%Y%m%d')
    projects = load_projects()
    all_users = get_all_users()
    
    # 获取当前用户的日报数据
    user_report = load_daily_report(username, today)
    
    log_activity(username, 'GET /daily_report')
    return render_template('daily_report.html', username=username, projects=projects, today=today, user_report=user_report, all_users=all_users)

@app.route('/submit_report', methods=['POST'])
def submit_report():
    if 'username' not in session:
        return redirect(url_for('index'))
    
    username = session['username']
    date = request.form.get('date')
    
    # 加载项目列表用于校验
    valid_projects = load_projects()
    
    # 收集多个事务
    transactions = []
    i = 0
    while f'work_content_{i}' in request.form:
        work_content = request.form.get(f'work_content_{i}')
        project = request.form.get(f'project_{i}')
        is_suspended = request.form.get(f'is_suspended_{i}') == 'on'
        suspended_reason = request.form.get(f'suspended_reason_{i}', '')
        suspended_end_date = request.form.get(f'suspended_end_date_{i}', '')
        is_help = request.form.get(f'is_help_{i}') == 'on'
        help_content = request.form.get(f'help_content_{i}', '')
        next_responsible = request.form.get(f'next_responsible_{i}', username)
        is_completed = request.form.get(f'is_completed_{i}') == 'on'
        
        # 校验项目名：如果填写了项目名，则必须在项目列表中
        if project and project.strip() not in valid_projects:
            return redirect(url_for('daily_report', error=f'事务{i+1}的项目名"{project}"不在项目列表中，请选择有效的项目名！'))
        
        if work_content.strip():
            transaction = {
                'work_content': work_content,
                'project': project,
                'is_suspended': is_suspended,
                'suspended_reason': suspended_reason,
                'suspended_end_date': suspended_end_date,
                'is_help': is_help,
                'help_content': help_content,
                'next_responsible': next_responsible,
                'is_completed': is_completed
            }
            transactions.append(transaction)
        
        i += 1
    
    if save_daily_report(username, date, transactions):
        log_activity(username, f'SUBMIT REPORT {date}')
        return redirect(url_for('daily_report', success='日报提交成功！'))
    else:
        return redirect(url_for('daily_report', error='日报提交失败，请重试！'))

@app.route('/get_report_data', methods=['GET'])
def get_report_data():
    if 'username' not in session:
        return jsonify({'error': '未登录'}), 401
    
    username = session['username']
    date = request.args.get('date')
    
    if not date:
        date = datetime.datetime.now().strftime('%Y%m%d')
    
    # 计算d-1日期
    date_obj = datetime.datetime.strptime(date, '%Y%m%d')
    prev_date = (date_obj - datetime.timedelta(days=1)).strftime('%Y%m%d')
    
    report_data = {}
    all_users = get_all_users()
    
    for user in all_users:
        # 获取d日工作
        today_works = load_daily_report(user, date)
        today_content = []
        for idx, work in enumerate(today_works, 1):
            # 构建完整的工作内容，包含所有信息
            content_str = f"{idx}. {work['work_content']}"
            
            # 添加项目信息
            if work.get('project'):
                content_str += f" [{work['project']}]"
            
            # 添加挂起信息
            if work.get('is_suspended'):
                if work.get('suspended_reason'):
                    content_str += f"\n   挂起原因: {work['suspended_reason']}"
                if work.get('suspended_end_date'):
                    content_str += f"\n   挂起结束日期: {work['suspended_end_date']}"
            
            # 添加求助信息
            if work.get('is_help'):
                if work.get('help_content'):
                    content_str += f"\n   求助内容: {work['help_content']}"
            
            # 添加负责人信息
            if work.get('next_responsible') and work['next_responsible'] != user:
                content_str += f"\n   负责人: {work['next_responsible']}"
            
            today_content.append(content_str)
        
        # 获取d-1日工作
        prev_works = load_daily_report(user, prev_date)
        prev_content = []
        for idx, work in enumerate(prev_works, 1):
            # 构建完整的工作内容，包含所有信息
            content_str = f"{idx}. {work['work_content']}"
            
            # 添加项目信息
            if work.get('project'):
                content_str += f" [{work['project']}]"
            
            # 添加挂起信息
            if work.get('is_suspended'):
                if work.get('suspended_reason'):
                    content_str += f"\n   挂起原因: {work['suspended_reason']}"
                if work.get('suspended_end_date'):
                    content_str += f"\n   挂起结束日期: {work['suspended_end_date']}"
            
            # 添加求助信息
            if work.get('is_help'):
                if work.get('help_content'):
                    content_str += f"\n   求助内容: {work['help_content']}"
            
            # 添加负责人信息
            if work.get('next_responsible') and work['next_responsible'] != user:
                content_str += f"\n   负责人: {work['next_responsible']}"
            
            prev_content.append(content_str)
        
        report_data[user] = {
            'd': '\n'.join(today_content),
            'd-1': '\n'.join(prev_content)
        }
    
    log_activity(username, f'GET REPORT DATA {date}')
    return jsonify(report_data)

@app.route('/download_docx', methods=['GET'])
def download_docx():
    if 'username' not in session:
        return redirect(url_for('index'))
    
    username = session['username']
    date = request.args.get('date')
    
    if not date:
        date = datetime.datetime.now().strftime('%Y%m%d')
    
    # 计算d-1日期
    date_obj = datetime.datetime.strptime(date, '%Y%m%d')
    prev_date = (date_obj - datetime.timedelta(days=1)).strftime('%Y%m%d')
    
    report_data = {}
    all_users = get_all_users()
    
    for user in all_users:
        # 获取d日工作
        today_works = load_daily_report(user, date)
        today_content = []
        for idx, work in enumerate(today_works, 1):
            # 构建完整的工作内容，包含所有信息
            content_str = f"{idx}. {work['work_content']}"
            
            # 添加项目信息
            if work.get('project'):
                content_str += f" [{work['project']}]"
            
            # 添加状态信息
            status = []
            if work.get('is_completed'):
                status.append("已完成")
            if work.get('is_suspended'):
                status.append("暂停中")
                if work.get('suspended_reason'):
                    content_str += f"\n   暂停原因: {work['suspended_reason']}"
                if work.get('suspended_end_date'):
                    content_str += f"\n   暂停截止: {work['suspended_end_date']}"
            if work.get('is_help'):
                status.append("需要帮助")
                if work.get('help_content'):
                    content_str += f"\n   帮助内容: {work['help_content']}"
            
            # 添加负责人信息
            if work.get('next_responsible') and work['next_responsible'] != user:
                content_str += f"\n   负责人: {work['next_responsible']}"
            
            if status:
                content_str += f" [{'/'.join(status)}]"
            
            today_content.append(content_str)
        
        # 获取d-1日工作
        prev_works = load_daily_report(user, prev_date)
        prev_content = []
        for idx, work in enumerate(prev_works, 1):
            # 构建完整的工作内容，包含所有信息
            content_str = f"{idx}. {work['work_content']}"
            
            # 添加项目信息
            if work.get('project'):
                content_str += f" [{work['project']}]"
            
            # 添加状态信息
            status = []
            if work.get('is_completed'):
                status.append("已完成")
            if work.get('is_suspended'):
                status.append("暂停中")
                if work.get('suspended_reason'):
                    content_str += f"\n   暂停原因: {work['suspended_reason']}"
                if work.get('suspended_end_date'):
                    content_str += f"\n   暂停截止: {work['suspended_end_date']}"
            if work.get('is_help'):
                status.append("需要帮助")
                if work.get('help_content'):
                    content_str += f"\n   帮助内容: {work['help_content']}"
            
            # 添加负责人信息
            if work.get('next_responsible') and work['next_responsible'] != user:
                content_str += f"\n   负责人: {work['next_responsible']}"
            
            if status:
                content_str += f" [{'/'.join(status)}]"
            
            prev_content.append(content_str)
        
        report_data[user] = {
            'd': '\n'.join(today_content) if today_content else '无',
            'd-1': '\n'.join(prev_content) if prev_content else '无'
        }
    
    # 生成docx文件
    filename = generate_docx_report(date, report_data)
    
    log_activity(username, f'DOWNLOAD DOCX {date}')
    
    # 使用send_file直接发送文件
    return send_file(filename, as_attachment=True, download_name=filename)


@app.route('/summary', methods=['POST'])
def summary():
    """获取指定日期所有组员工作内容"""
    data = request.json
    username = data.get('username')
    password = data.get('password')
    date = data.get('date')
    
    # 验证用户名密码
    passwords = load_passwords()
    if username not in passwords or passwords[username] != password:
        return jsonify({'error': '用户名或密码错误'}), 401
    
    if not date:
        return jsonify({'error': '日期不能为空'}), 400
    
    # 收集所有用户的日报数据
    all_data = {}
    all_users = get_all_users()
    
    for user in all_users:
        user_report = load_daily_report(user, date)
        all_data[user] = user_report
    
    log_activity(username, f'API SUMMARY {date}')
    return jsonify(all_data)

@app.route('/upload', methods=['POST'])
def upload():
    """上传项目名称列表"""
    data = request.json
    username = data.get('username')
    password = data.get('password')
    projects = data.get('projects', [])
    
    # 验证用户名密码
    passwords = load_passwords()
    if username not in passwords or passwords[username] != password:
        return jsonify({'error': '用户名或密码错误'}), 401
    
    if not isinstance(projects, list):
        return jsonify({'error': '项目列表格式错误'}), 400
    
    if save_projects(projects):
        log_activity(username, f'API UPLOAD PROJECTS: {len(projects)} items')
        return jsonify({'success': True, 'message': '项目列表上传成功'})
    else:
        return jsonify({'error': '项目列表保存失败'}), 500

@app.route('/search_projects', methods=['GET'])
def api_search_projects():
    """搜索项目名称"""
    if 'username' not in session:
        # 验证是否通过API调用（带有用户名密码参数）
        username = request.args.get('username')
        password = request.args.get('password')
        passwords = load_passwords()
        if not (username in passwords and passwords[username] == password):
            return jsonify({'error': '未授权'}), 401
    
    keyword = request.args.get('keyword', '')
    results = search_projects(keyword)
    
    return jsonify(results)

# 静态文件路由
@app.route('/static/<path:filename>')
def static_file(filename):
    return app.send_static_file(filename)

if __name__ == '__main__':
    # 确保必要的目录存在
    ensure_directories()
    # 创建session目录
    if not os.path.exists(app.config['SESSION_FILE_DIR']):
        os.makedirs(app.config['SESSION_FILE_DIR'])
    
    # 启动Flask应用
    # 检查是否在Windows环境下运行
    import platform
    is_windows = platform.system() == 'Windows'
    
    # 证书路径
    if is_windows:
        # Windows环境下的证书路径示例
        cert_path = 'certs/cert.pem'
        key_path = 'certs/privkey.pem'
        chain_path = 'certs/chain.pem'
        combined_cert_path = 'certs/combined_cert.pem'
    else:
        # Linux环境下的证书路径
        # （已屏蔽需要配置！）
        cert_path = ''
        key_path = ''
        chain_path = ''
        combined_cert_path = '/tmp/combined_cert.pem'
    
    # 检查证书文件是否存在
    cert_exists = os.path.exists(cert_path)
    key_exists = os.path.exists(key_path)
    chain_exists = os.path.exists(chain_path)
    
    print("启动Flask日报管理系统...")
    
    if cert_exists and key_exists:
        # 证书和私钥都存在，尝试使用HTTPS
        print("检测到证书文件，尝试启动HTTPS服务...")
        print(f"访问地址: https://localhost:5000")
        
        # 尝试合并证书
        try:
            # 创建证书目录（如果在Windows下且目录不存在）
            if is_windows and not os.path.exists('certs'):
                os.makedirs('certs')
                print("已创建证书目录，请将证书文件放入certs文件夹")
                # 如果刚创建目录，证书可能还没放进去，使用HTTP模式
                cert_exists = False
                key_exists = False
            
            if cert_exists and key_exists:
                # 读取证书内容
                with open(cert_path, 'r') as cert_file:
                    cert_content = cert_file.read()
                
                # 合并中间证书（如果存在）
                combined_content = cert_content
                if chain_exists:
                    with open(chain_path, 'r') as chain_file:
                        chain_content = chain_file.read()
                    combined_content += '\n' + chain_content
                    print("已包含中间证书")
                
                # 写入合并后的证书
                with open(combined_cert_path, 'w') as combined_file:
                    combined_file.write(combined_content)
                
                print(f"已创建组合证书文件: {combined_cert_path}")
                
                # 使用HTTPS启动
                app.run(debug=False, host='0.0.0.0', port=5000,
                        ssl_context=(combined_cert_path, key_path))
        except Exception as e:
            print(f"HTTPS配置失败: {e}")
            print("注意：在Windows环境下，证书路径应为certs文件夹")
            print("请将证书文件复制到以下位置:")
            print(f"  - 证书: {cert_path}")
            print(f"  - 私钥: {key_path}")
            print(f"  - 中间证书: {chain_path} (可选)")
    
    # 如果证书不存在或HTTPS配置失败，使用HTTP模式启动（开发模式）
    print("\n使用HTTP模式启动（开发模式）...")
    print("访问地址: http://localhost:5000")
    print("注意：生产环境请确保证书配置正确以使用HTTPS")
    app.run(debug=True, host='0.0.0.0', port=5000)